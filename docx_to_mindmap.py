#!/usr/bin/env python3
"""
PDF 转思维导图工具
读取 PDF 文件，提取内容并生成 Markdown 格式的思维导图
"""

import os
import re
from pathlib import Path
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv


def extract_docx_text(docx_path):
    """提取 Word 文档文本内容"""
    import sys

    print(f"正在读取 DOCX: {docx_path}")
    sys.stdout.flush()

    doc = Document(docx_path)
    text_content = []

    # 读取所有段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            text_content.append(text)

        if (i + 1) % 100 == 0:
            print(f"处理第 {i+1} 段")
            sys.stdout.flush()

    # 读取所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    text_content.append(text)

    full_text = "\n".join(text_content)
    print(f"DOCX 文本提取完成")
    print(f"总段落数: {len(doc.paragraphs)}")
    print(f"总字符数: {len(full_text)}")
    sys.stdout.flush()

    return full_text


def extract_pdf_text_pymupdf(pdf_path, max_pages=None):
    """使用 PyMuPDF 提取 PDF 文本内容"""
    import sys
    text_content = []

    print(f"正在读取 PDF (PyMuPDF): {pdf_path}")
    sys.stdout.flush()

    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    pages_to_process = min(max_pages, total_pages) if max_pages else total_pages

    print(f"PDF 总页数: {total_pages}，将处理: {pages_to_process} 页")
    sys.stdout.flush()

    for i in range(pages_to_process):
        if (i + 1) % 10 == 0:  # 每10页输出一次
            print(f"处理第 {i+1}/{pages_to_process} 页")
            sys.stdout.flush()
        page = doc[i]
        text = page.get_text()
        if text and text.strip():
            text_content.append(text)

    doc.close()

    print(f"PDF 文本提取完成，共 {len(text_content)} 页有效内容")
    print(f"总字符数: {sum(len(t) for t in text_content)}")
    sys.stdout.flush()
    return "\n\n".join(text_content)


def extract_pdf_text(pdf_path, max_pages=None):
    """提取 PDF 文本内容（优先使用 PyMuPDF）"""
    import sys

    # 先尝试使用 PyMuPDF
    try:
        text = extract_pdf_text_pymupdf(pdf_path, max_pages)
        if text and len(text.strip()) > 100:  # 如果提取到了有效文本
            return text
        print(f"PyMuPDF 提取文本太少，尝试使用 pdfplumber")
        sys.stdout.flush()
    except Exception as e:
        print(f"PyMuPDF 提取失败: {e}，尝试使用 pdfplumber")
        sys.stdout.flush()

    # 如果 PyMuPDF 失败，尝试 pdfplumber
    text_content = []
    print(f"正在读取 PDF (pdfplumber): {pdf_path}")
    sys.stdout.flush()

    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        pages_to_process = min(max_pages, total_pages) if max_pages else total_pages

        print(f"PDF 总页数: {total_pages}，将处理: {pages_to_process} 页")
        sys.stdout.flush()

        for i, page in enumerate(pdf.pages[:pages_to_process]):
            if (i + 1) % 10 == 0:  # 每10页输出一次
                print(f"处理第 {i+1}/{pages_to_process} 页")
                sys.stdout.flush()
            text = page.extract_text()
            if text:
                text_content.append(text)

    print(f"PDF 文本提取完成，共 {len(text_content)} 页")
    sys.stdout.flush()
    return "\n\n".join(text_content)


def generate_mindmap_md(pdf_text, pdf_name):
    """使用 AI API 生成思维导图 Markdown"""
    import sys

    api_key = os.getenv("DASHSCOPE_API_KEY")

    if not api_key:
        print("警告: 未找到 DASHSCOPE_API_KEY，将使用简单的结构提取")
        sys.stdout.flush()
        return generate_simple_mindmap(pdf_text, pdf_name)

    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    )

    # 如果文本太长，分批处理
    text_length = len(pdf_text)
    max_chunk_size = 100000  # 每次处理最多 10万字符

    print(f"PDF 文本总长度: {text_length} 字符")
    sys.stdout.flush()

    prompt = f"""请仔细分析以下教材内容，生成一个详细完整的 Markdown 格式背诵脑图。

【格式要求 - 非常重要】
1. 必须使用标准 Markdown 标题格式：
   - 一级标题用 # （书名）
   - 二级标题用 ## （章节名）
   - 三级标题用 ### （小节名）
   - 四级标题用 #### （知识点）
   - 五级标题用 ##### （子知识点）

2. 示例格式：
```
# 教材名称

## 第一章 章节名称

### 1.1 小节名称

#### 知识点1 3
（表示这个知识点有3条要点）

##### 子知识点1.1 2

##### 子知识点1.2

#### 知识点2 5

### 1.2 小节名称

#### 知识点3 4
```

【内容要求】
1. **完整性**：必须覆盖文本中的所有章节、所有知识点，不能遗漏
2. **详细性**：每个知识点都要列出来，包括：
   - 概念定义
   - 分类类型
   - 特点特征
   - 过程步骤
   - 影响因素
   - 应用意义
3. **层次性**：
   - 章节 → 节 → 知识点 → 子知识点，层次清晰
   - 每个层级都要完整
4. **标注数量规则**（非常重要，请严格遵守）：
   - **最外层叶子节点**（没有子节点的节点）：**必须标注数字**，表示该知识点有几条要背诵的内容
     * 示例：##### 溶解性 3  （表示溶解性有3条要背诵的内容）
   - **中间节点**（有子节点的节点）：**默认不标注数字**
     * 标注数字的唯一例外：当且仅当原文中有明确的条数标注时才标数字
     * 明确的条数标注包括：
       - 序号标注：(1)(2)(3)、（1）（2）（3）、①②③、1、2、3、一、二、三等
       - 明确数量词：如"有4种化学性质"、"包括3个方面"、"分为5类"等
     * 如果原文只是普通叙述，没有明显的序号或数量词，则中间节点不标数字
   - 判断标准：
     * 如果这个标题下面还有子标题（###、####、#####等），就是中间节点
       → 检查原文是否有明确的序号或数量词，有则标数字，无则不标
     * 如果这个标题下面没有子标题了，就是最外层叶子节点
       → 必须标数字
   - 完整示例：
     ```
     ## 第一章 氨基酸              ← 不标数字（下面有子节点）
     ### 1.1 氨基酸的性质          ← 不标数字（下面有子节点）
     #### 物理性质                ← 不标数字（下面有子节点）
     ##### 溶解性 3                ← 标数字（最外层叶子节点，下面没有子节点了）
     ##### 熔点 2                  ← 标数字（最外层叶子节点）
     ##### 颜色 1                  ← 标数字（最外层叶子节点）
     #### 化学性质 4               ← 可标数字（原文明确写了"4种化学性质"）
     ##### 酸碱性 2                ← 标数字（最外层叶子节点）
     ##### 成盐反应 3              ← 标数字（最外层叶子节点）
     ##### 氧化还原反应 2          ← 标数字（最外层叶子节点）
     ##### 缩合反应 4              ← 标数字（最外层叶子节点）
     ```
5. **不写答案**：只列出知识点名称和要点数量，不写具体内容

【分析策略】
1. 先识别所有章节标题
2. 再识别每章的所有小节
3. 然后列出每节的所有知识点
4. 最后补充每个知识点的子知识点

书名：{pdf_name}

教材内容：
{pdf_text[:max_chunk_size]}

请严格按照上述格式生成完整详细的思维导图："""

    print(f"正在使用 AI 生成详细思维导图...")
    sys.stdout.flush()

    completion = client.chat.completions.create(
        model="qwen3-max",
        messages=[
            {
                "role": "system",
                "content": "你是一个专业的教育专家，擅长分析教材内容并生成结构化的背诵脑图。你必须确保生成的内容详细完整，不遗漏任何知识点。"
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.3,  # 降低温度以获得更稳定的输出
    )

    result = completion.choices[0].message.content
    print(f"思维导图生成完成，长度: {len(result)} 字符")
    sys.stdout.flush()

    return result


def generate_simple_mindmap(pdf_text, pdf_name):
    """简单的思维导图生成（不使用 API）"""
    lines = pdf_text.split('\n')

    # 简单的结构提取逻辑
    mindmap = [f"# {pdf_name}"]
    mindmap.append("")

    current_section = None
    for line in lines[:500]:  # 只处理前 500 行
        line = line.strip()
        if not line:
            continue

        # 检测章节标题（通常是数字开头或者较短的行）
        if re.match(r'^第[一二三四五六七八九十\d]+[章节]', line) or \
           re.match(r'^\d+[\.\、]', line):
            if len(line) < 50:
                mindmap.append(f"## {line}")
                current_section = line
        elif current_section and len(line) < 100:
            mindmap.append(f"- {line}")

    return "\n".join(mindmap)


def convert_md_to_xmind(md_file, output_dir="output"):
    """使用 md2xmind 转换 Markdown 到 XMind"""
    import md2xmind

    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)

    xmind_file = output_path / f"{Path(md_file).stem}.xmind"

    print(f"正在转换 {md_file} 到 XMind 格式...")

    try:
        # 使用 md2xmind API
        topic_name = Path(md_file).stem
        md2xmind.start_trans_file(str(md_file), str(xmind_file), topic_name)
        print(f"转换成功: {xmind_file}")
        return xmind_file
    except Exception as e:
        print(f"转换失败: {e}")
        import traceback
        traceback.print_exc()
        return None


def main():
    """主函数"""
    # 加载环境变量
    load_dotenv()

    data_dir = Path("data")
    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)

    # 获取所有 DOCX 文件
    docx_files = list(data_dir.glob("*.docx"))

    if not docx_files:
        print("未找到 DOCX 文件")
        return

    print(f"找到 {len(docx_files)} 个 DOCX 文件")

    for docx_file in docx_files:
        print(f"\n{'='*60}")
        print(f"处理: {docx_file.name}")
        print(f"{'='*60}")

        try:
            # 1. 提取 DOCX 文本
            doc_text = extract_docx_text(docx_file)

            if not doc_text or len(doc_text.strip()) == 0:
                print(f"警告: DOCX 文本为空或无法提取，跳过此文件")
                continue

            # 2. 生成思维导图 Markdown
            doc_name = docx_file.stem
            mindmap_md = generate_mindmap_md(doc_text, doc_name)

            # 3. 保存 Markdown 文件
            md_file = output_dir / f"{doc_name}.md"
            with open(md_file, 'w', encoding='utf-8') as f:
                f.write(mindmap_md)
            print(f"已保存 Markdown: {md_file}")

            # 4. 转换为 XMind
            xmind_file = convert_md_to_xmind(md_file, output_dir)

        except Exception as e:
            print(f"处理 {docx_file.name} 时出错: {e}")
            import traceback
            traceback.print_exc()
            continue

    print(f"\n{'='*60}")
    print("全部处理完成！")
    print(f"输出目录: {output_dir.absolute()}")


if __name__ == "__main__":
    main()
